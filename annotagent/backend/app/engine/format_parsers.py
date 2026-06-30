"""Format-specific parsers for CodebookAgent's Ingestor role.

Each parser returns `IngestResult` with cleaned text, detected tables, optional
analysis-friendly rows (for annotator-sheet XLSX), and structured warnings.
All parsers are defensive: bounded by time + size + row/char caps; every
exception is caught and surfaced as a warning instead of crashing the server.
"""
from __future__ import annotations

import asyncio
import csv
import io
import json
import logging
from dataclasses import dataclass, field
from typing import Any

logger = logging.getLogger(__name__)

# Defense-in-depth caps
MAX_TEXT_CHARS = 100_000       # Drafter context cap
MAX_TABLE_ROWS = 5_000         # per-table safety
PER_PARSER_TIMEOUT_S = 30.0    # hard timeout per parser


@dataclass
class Table:
    """A parsed table. May carry label annotations if source is annotator data."""
    name: str                                   # sheet name or "body"
    header: list[str] = field(default_factory=list)
    rows: list[dict[str, Any]] = field(default_factory=list)


@dataclass
class IngestResult:
    """What every format parser returns. Never raises — errors surface as warnings."""
    clean_text: str = ""
    tables: list[Table] = field(default_factory=list)
    # If input looks like annotator data, we can emit analysis-friendly rows here.
    # Each row: {row_id, user_id, timestamp, sentence, behavior_side, annotations, source}
    analysis_rows: list[dict[str, Any]] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    ok: bool = True


def _cap_text(s: str) -> tuple[str, str | None]:
    if len(s) > MAX_TEXT_CHARS:
        return s[:MAX_TEXT_CHARS], f"Text truncated to first {MAX_TEXT_CHARS:,} chars (input was {len(s):,})."
    return s, None


async def parse_file(data: bytes, filename: str, mime: str = "", only_sheet: str | None = None) -> IngestResult:
    """Top-level dispatch. Returns IngestResult; never raises.

    Detection order: extension → MIME → first 8 bytes.
    ``only_sheet`` restricts XLSX parsing to a single named sheet (else all sheets
    are merged, the default).
    """
    ext = (filename.rsplit(".", 1)[-1] or "").lower() if "." in filename else ""
    mime_lc = (mime or "").lower()

    async def _run(parser, *args):
        try:
            return await asyncio.wait_for(parser(*args), timeout=PER_PARSER_TIMEOUT_S)
        except asyncio.TimeoutError:
            return IngestResult(warnings=[f"Parser timed out after {PER_PARSER_TIMEOUT_S}s"], ok=False)
        except Exception as e:
            logger.warning(f"Parser {parser.__name__} failed on {filename}: {e}")
            return IngestResult(warnings=[f"{parser.__name__} failed: {type(e).__name__}: {e}"], ok=False)

    # Extension takes priority; MIME is often generic octet-stream from FastAPI
    if ext in ("xlsx", "xlsm"):
        return await _run(_parse_xlsx, data, filename, only_sheet)
    if ext in ("docx",):
        return await _run(_parse_docx, data, filename)
    if ext == "pdf":
        return await _run(_parse_pdf, data, filename)
    if ext == "csv":
        return await _run(_parse_csv, data, filename)
    if ext == "json":
        return await _run(_parse_json, data, filename)
    if ext in ("txt", "md"):
        return await _run(_parse_text, data, filename)

    # MIME fallback
    if "pdf" in mime_lc:
        return await _run(_parse_pdf, data, filename)
    if "spreadsheet" in mime_lc or "excel" in mime_lc:
        return await _run(_parse_xlsx, data, filename, only_sheet)
    if "word" in mime_lc or "officedocument.wordprocessingml" in mime_lc:
        return await _run(_parse_docx, data, filename)
    if "csv" in mime_lc:
        return await _run(_parse_csv, data, filename)
    if "json" in mime_lc:
        return await _run(_parse_json, data, filename)
    if "text/" in mime_lc:
        return await _run(_parse_text, data, filename)

    return IngestResult(
        warnings=[f"Unknown file type (ext={ext!r}, mime={mime!r}). Supported: pdf docx xlsx csv json txt."],
        ok=False,
    )


# ─── individual parsers ──────────────────────────────────────────────────────

async def _parse_text(data: bytes, filename: str) -> IngestResult:
    text = data.decode("utf-8", errors="replace")
    capped, warn = _cap_text(text)
    w = [warn] if warn else []
    return IngestResult(clean_text=capped, warnings=w)


async def _parse_csv(data: bytes, filename: str) -> IngestResult:
    text = data.decode("utf-8-sig", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    rows = []
    header = reader.fieldnames or []
    warnings: list[str] = []
    for i, row in enumerate(reader):
        if i >= MAX_TABLE_ROWS:
            warnings.append(f"CSV truncated at {MAX_TABLE_ROWS} rows")
            break
        rows.append(row)
    table = Table(name="body", header=header, rows=rows)
    # Also produce a text summary for the Drafter
    preview = [", ".join(header)]
    for r in rows[:40]:
        preview.append(", ".join(str(v or "") for v in r.values()))
    capped, warn = _cap_text("\n".join(preview))
    if warn:
        warnings.append(warn)
    return IngestResult(clean_text=capped, tables=[table], warnings=warnings)


async def _parse_json(data: bytes, filename: str) -> IngestResult:
    text = data.decode("utf-8", errors="replace")
    try:
        obj = json.loads(text)
    except json.JSONDecodeError as e:
        return IngestResult(warnings=[f"Invalid JSON: {e}"], ok=False)

    # If it looks like a codebook already, return the raw text (Drafter will pass-through)
    warnings: list[str] = []
    if isinstance(obj, dict) and "dimensions" in obj and isinstance(obj["dimensions"], list):
        warnings.append("Input JSON looks like an existing codebook — will be validated directly.")

    capped, warn = _cap_text(json.dumps(obj, indent=2, ensure_ascii=False))
    if warn:
        warnings.append(warn)
    return IngestResult(clean_text=capped, warnings=warnings)


async def _parse_pdf(data: bytes, filename: str) -> IngestResult:
    try:
        from pypdf import PdfReader
    except ImportError:
        return IngestResult(warnings=["pypdf not installed — cannot parse PDF"], ok=False)

    reader = PdfReader(io.BytesIO(data))
    chunks = []
    warnings: list[str] = []
    total_chars = 0
    for i, page in enumerate(reader.pages):
        if total_chars > MAX_TEXT_CHARS:
            warnings.append(f"PDF truncated after page {i}")
            break
        try:
            page_text = page.extract_text() or ""
        except Exception as e:
            warnings.append(f"Page {i+1} extraction failed: {e}")
            continue
        chunks.append(page_text)
        total_chars += len(page_text)

    text = "\n\n".join(chunks).strip()
    if not text:
        warnings.append(
            "No extractable text — PDF may be scanned/image-only. "
            "Paste the codebook text in Door B instead."
        )
        return IngestResult(clean_text="", warnings=warnings, ok=False)

    capped, warn = _cap_text(text)
    if warn:
        warnings.append(warn)
    return IngestResult(clean_text=capped, warnings=warnings)


async def _parse_docx(data: bytes, filename: str) -> IngestResult:
    try:
        import docx
    except ImportError:
        return IngestResult(warnings=["python-docx not installed — cannot parse DOCX"], ok=False)

    doc = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Tables often carry the codebook — extract them as structured data
    tables: list[Table] = []
    for t_idx, t in enumerate(doc.tables):
        rows = t.rows
        if not rows:
            continue
        header = [c.text.strip() for c in rows[0].cells]
        data_rows = []
        for r in rows[1:MAX_TABLE_ROWS + 1]:
            row = {h: c.text.strip() for h, c in zip(header, r.cells)}
            data_rows.append(row)
        tables.append(Table(name=f"table_{t_idx}", header=header, rows=data_rows))

    # Build clean_text = paragraphs + flattened tables
    parts = paragraphs[:]
    for tbl in tables:
        parts.append(f"\n[table: {tbl.name}]")
        parts.append(" | ".join(tbl.header))
        for r in tbl.rows[:30]:
            parts.append(" | ".join(str(r.get(h, "")) for h in tbl.header))

    text = "\n".join(parts)
    capped, warn = _cap_text(text)
    warnings = [warn] if warn else []
    if not text.strip():
        warnings.append("DOCX contained no extractable text.")
        return IngestResult(warnings=warnings, ok=False)
    return IngestResult(clean_text=capped, tables=tables, warnings=warnings)


def xlsx_content_sheets(data: bytes) -> list[str]:
    """Rule-based list of XLSX sheets that carry tabular content (>= 2 non-empty
    rows). Used at upload to decide whether to ask about multi-sheet merging.
    Returns [] if the file is unreadable as XLSX."""
    try:
        import openpyxl
    except ImportError:
        return []
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    except Exception:
        return []
    out: list[str] = []
    for name in wb.sheetnames:
        ws = wb[name]
        nonempty = 0
        for row in ws.iter_rows(values_only=True):
            if any(c is not None and str(c).strip() for c in row):
                nonempty += 1
                if nonempty >= 2:
                    out.append(name)
                    break
    return out


async def _parse_xlsx(data: bytes, filename: str, only_sheet: str | None = None) -> IngestResult:
    try:
        import openpyxl
    except ImportError:
        return IngestResult(warnings=["openpyxl not installed"], ok=False)

    # read_only + data_only for memory safety on large files
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    warnings: list[str] = []
    tables: list[Table] = []
    analysis_rows: list[dict[str, Any]] = []

    sheet_names = wb.sheetnames
    if only_sheet is not None:
        if only_sheet not in sheet_names:
            return IngestResult(warnings=[f"Sheet {only_sheet!r} not found in workbook."], ok=False)
        sheet_names = [only_sheet]

    for sheet_name in sheet_names:
        ws = wb[sheet_name]
        # Collect first ~MAX_TABLE_ROWS rows; detect header
        raw_rows: list[tuple] = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= MAX_TABLE_ROWS + 5:
                warnings.append(f"Sheet '{sheet_name}' truncated at {MAX_TABLE_ROWS} rows")
                break
            raw_rows.append(row)

        if not raw_rows:
            continue

        # Detect annotator-style header (meta row in r1, actual header in r2)
        r1 = [str(c or "").strip() for c in raw_rows[0]] if raw_rows else []
        r2 = [str(c or "").strip() for c in raw_rows[1]] if len(raw_rows) > 1 else []
        if r1[:3] == ["#", "User", "FL"] or r1[:3] == ["#", "User", "CW"] or (
            len(r1) >= 3 and r1[0] == "#" and r1[1].startswith("User")
        ):
            header = r2
            data_start = 2
            is_annotator = True
            annotator_name = r1[2] if len(r1) >= 3 else ""
            warnings.append(
                f"Sheet '{sheet_name}' detected as annotator sheet "
                f"(annotator={annotator_name!r}); header taken from row 2."
            )
        else:
            header = r1
            data_start = 1
            is_annotator = False

        # Forward-fill continuation rows + canonicalize.
        # Annotator sheets fill each column independently (a single Coding-theme
        # column legitimately spans many level rows). Taxonomy sheets are nested
        # (Type -> Function -> Code -> Subcode, left-to-right): when a parent
        # column gets a new value, deeper (right-of) columns must reset, else a
        # stale child cell attaches to the wrong parent.
        last: list[Any] = [None] * len(header)
        rows: list[dict[str, Any]] = []
        for row_tuple in raw_rows[data_start:]:
            if not any(v is not None for v in row_tuple):
                continue
            row_dict: dict[str, Any] = {}
            for i, h in enumerate(header):
                v = row_tuple[i] if i < len(row_tuple) else None
                if v is None or (isinstance(v, str) and not v.strip()):
                    row_dict[h] = last[i]
                else:
                    val = v.strip() if isinstance(v, str) else v
                    row_dict[h] = val
                    last[i] = val
                    if not is_annotator:
                        for j in range(i + 1, len(header)):
                            last[j] = None
            rows.append(row_dict)

        tables.append(Table(name=sheet_name, header=header, rows=rows))

        # If this sheet is annotator data, extract analysis rows
        if _looks_like_annotator_sheet(header):
            analysis_rows.extend(_extract_analysis_rows(sheet_name, header, rows))

    # Build clean_text = header + first N rows per sheet, so Drafter can see structure
    parts = []
    for t in tables:
        parts.append(f"\n[sheet: {t.name}]")
        parts.append(" | ".join(h or "" for h in t.header))
        for r in t.rows[:25]:
            parts.append(" | ".join(str(r.get(h, "") or "") for h in t.header))

    capped, warn = _cap_text("\n".join(parts))
    if warn:
        warnings.append(warn)
    return IngestResult(
        clean_text=capped,
        tables=tables,
        analysis_rows=analysis_rows,
        warnings=warnings,
    )


# ─── annotator-sheet detection & analysis-row extraction ────────────────────

def _looks_like_annotator_sheet(header: list[str]) -> bool:
    lower = [h.lower() for h in header if h]
    has_coding = any("coding theme" in h or "theme" == h for h in lower)
    has_quote = any("quote" in h for h in lower) or any("sentence" in h for h in lower)
    has_level = any(h in ("level", "subcategory") for h in lower)
    return has_coding and (has_quote or has_level)


def _canon_label(v: Any) -> str:
    if v is None:
        return ""
    s = str(v).strip()
    if not s:
        return ""
    # Normalize "Yes, it's a confession" -> "Yes"
    lower = s.lower()
    if lower.startswith("yes"):
        return "Yes"
    if lower.startswith("no"):
        return "No"
    # Strip redundant "layer"/"level" suffix
    for suf in (" layer", " level"):
        if lower.endswith(suf):
            return s[: -len(suf)].strip().title()
    return s


def _canon_multi(v: Any) -> list[str]:
    """Split "A & B & C" into a cleaned label list."""
    if v is None:
        return []
    parts = [p.strip() for p in str(v).split("&")]
    return [_canon_label(p) for p in parts if p.strip()]


def _extract_analysis_rows(sheet_name: str, header: list[str], rows: list[dict]) -> list[dict]:
    """Produce one analysis-friendly row per annotation record."""
    # Column resolver (case-insensitive)
    def find_col(*names: str) -> str | None:
        lower_map = {h.lower(): h for h in header if h}
        for n in names:
            if n.lower() in lower_map:
                return lower_map[n.lower()]
        return None

    col_quote = find_col("Relevant quotes", "Sentence", "Text")
    col_theme = find_col("Coding theme", "Theme")
    col_level = find_col("Level", "Subcategory")
    col_topic = find_col("Topic")
    col_topic_cat = find_col("Topic thematic category")
    col_time = find_col("Time stamp", "Timestamp")
    col_user = find_col("User", "user_id", "User ")
    col_behavior = find_col("Behavior (Self-disclosure/AI behavior)", "Behavior")

    out: list[dict] = []
    for i, r in enumerate(rows):
        sentence = r.get(col_quote) if col_quote else None
        if not sentence:
            continue
        dim = r.get(col_theme) if col_theme else None
        raw_level = r.get(col_level) if col_level else None
        # Multi-label detection
        if raw_level and "&" in str(raw_level):
            label = _canon_multi(raw_level)
        else:
            label = _canon_label(raw_level)

        annotations: dict[str, Any] = {}
        if dim:
            annotations[str(dim).strip()] = label
        if col_topic and r.get(col_topic):
            annotations["Topic"] = _canon_label(r[col_topic])
        if col_topic_cat and r.get(col_topic_cat):
            annotations["Topic thematic category"] = _canon_label(r[col_topic_cat])

        side = None
        if col_behavior and r.get(col_behavior):
            b = str(r[col_behavior]).lower()
            if "ai" in b:
                side = "ai"
            elif "self" in b:
                side = "user"

        out.append({
            "row_id": i + 1,
            "user_id": str(r.get(col_user, "")).strip() if col_user else "",
            "timestamp": str(r.get(col_time, "")).strip() if col_time else "",
            "sentence": str(sentence).strip(),
            "behavior_side": side,
            "annotations": annotations,
            "source": {"sheet": sheet_name, "xlsx_row": i + 1},
        })
    return out
